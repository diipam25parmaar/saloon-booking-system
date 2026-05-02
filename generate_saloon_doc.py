"""
generate_saloon_doc.py
Generates SaloonBookingSchedulerDocumentation.docx
Run: python generate_saloon_doc.py
Requires: pip install python-docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.size = Pt(24)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.color.rgb = RGBColor(0, 102, 204)
        run.font.size = Pt(18)
    elif level == 3:
        run.font.color.rgb = RGBColor(51, 51, 51)
        run.font.size = Pt(14)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def create_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)
    
    doc.add_paragraph()

def generate_document():
    doc = Document()
    
    # Title Page
    doc.add_picture = lambda *args, **kwargs: None # Ignore if no logo
    title = add_heading(doc, "Saloon Booking Scheduler", level=1)
    subtitle = add_paragraph(doc, "Technical Documentation & System Architecture")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].italic = True
    doc.add_page_break()
    
    # Overview
    add_heading(doc, "1. Project Overview", level=2)
    add_paragraph(doc, "The Saloon Booking Scheduler is a full-stack, production-ready appointment management platform built for modern salon operations. It features a role-based web application allowing administrators to manage salon services, working hours, and bookings, while customers can register, browse available slots, and confirm appointments via a self-service flow.")
    
    create_table(doc, ["Directory", "Purpose"], [
        ["saloon-booking-backend", "Laravel 12 REST API — auth, categories, services, bookings, slots, working-time rules"],
        ["saloon-booking-frontend", "React 19 + Vite SPA — admin panel and customer booking UI"]
    ])
    
    # Features
    add_heading(doc, "2. Features", level=2)
    add_heading(doc, "Admin Features", level=3)
    add_bullet(doc, "Secure admin login with role guard")
    add_bullet(doc, "Categories CRUD (create, edit, delete)")
    add_bullet(doc, "Services CRUD (create, edit, delete with duration, price, and category linking)")
    add_bullet(doc, "Working-Time Rules management")
    add_bullet(doc, "Booking Management with client details")
    add_bullet(doc, "Slot Overview by Date")
    
    add_heading(doc, "Customer Features", level=3)
    add_bullet(doc, "Self-registration and secure authentication")
    add_bullet(doc, "Category-first booking flow")
    add_bullet(doc, "Real-time slot grid generated on-demand")
    add_bullet(doc, "Auto-populates session user data")
    
    # Tech Stack
    add_heading(doc, "3. Tech Stack", level=2)
    create_table(doc, ["Tool", "Version", "Role"], [
        ["PHP", "8.2+", "Backend Runtime"],
        ["Laravel", "12.x", "Backend Framework"],
        ["Laravel Sanctum", "4.x", "Token Authentication"],
        ["SQLite/MySQL", "latest", "Database"],
        ["React", "19.x", "Frontend UI"],
        ["React Router", "7.x", "Routing"],
        ["Vite", "8.x", "Build tool"]
    ])
    
    # System Flow
    add_heading(doc, "4. Full System Flow", level=2)
    add_heading(doc, "Admin Flow", level=3)
    add_bullet(doc, "Navigate to /login and enter admin credentials")
    add_bullet(doc, "Token stored in localStorage, redirect to /admin")
    add_bullet(doc, "Manage Working-Time Rules, Categories, Services")
    add_bullet(doc, "View and cancel appointments in /admin/bookings")
    
    add_heading(doc, "Customer Flow", level=3)
    add_bullet(doc, "Register or Login to access /booking")
    add_bullet(doc, "Select Category -> Select Service -> Select Date")
    add_bullet(doc, "Slots load automatically via API, filtering booked and past slots")
    add_bullet(doc, "Confirm appointment and view success message")
    
    # API Overview
    add_heading(doc, "5. API Overview", level=2)
    add_paragraph(doc, "Base URL: /api")
    create_table(doc, ["Method", "Endpoint", "Auth", "Description"], [
        ["POST", "/login", "Public", "Authenticate user and return token"],
        ["GET", "/available-slots", "Public", "Get slots for date + service"],
        ["POST", "/appointments", "Bearer Token", "Create new appointment"],
        ["POST", "/services", "Admin Token", "Create new service"],
        ["GET", "/working-time-rules", "Admin Token", "List rules"]
    ])
    
    # Database Design
    add_heading(doc, "6. Database Design", level=2)
    create_table(doc, ["Table", "Description"], [
        ["users", "Stores admins and customers (email, password, role)"],
        ["categories", "Service categories"],
        ["services", "Available services linked to categories"],
        ["appointments", "Booked slots linked to user and service"],
        ["working_time_rules", "Availability rules (day of week or specific date)"]
    ])
    
    # Slot Logic
    add_heading(doc, "7. Slot Generation Logic", level=2)
    add_paragraph(doc, "Slots are generated on-demand by Loading the requested service duration, cross-referencing working-time rules for the requested date, and filtering out any time blocks that intersect with existing appointments in the database. Overlap detection uses strict interval logic: start_at < slotEnd AND end_at > slotStart.")
    
    output_path = os.path.abspath('SaloonBookingSchedulerDocumentation.docx')
    doc.save(output_path)
    
    # Fix the ascii printing issue that was encountered previously
    print("Documentation saved: " + output_path)

if __name__ == "__main__":
    generate_document()
